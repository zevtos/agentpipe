#!/usr/bin/env python3
"""Integration tests for the dash-sanitize coverage gaps (GAP 2 + GAP 3).

Unlike test_viz.py (pure-python chart objects), this exercises the FULL docx
pipeline and therefore REQUIRES python-docx. Run it with a python that has
python-docx installed (the skill venv created by ensure_env, which also carries
numpy+matplotlib):

    ~/.local/share/agentpipe/gost-report/venv/bin/python scripts/test_sanitize_integration.py

Mirrors the style of scripts/validate-*.py and scripts/test_viz.py:
assertion-based, runnable via `python3`, exits non-zero on any failure.

What it covers:
  GAP 2 — Report.table() sanitizes every header+data cell. Build a real Report,
          add a table whose cells contain em/en-dashes plus a text() paragraph
          with an em-dash, save to a temp .docx, reopen it, and assert NO
          prohibited [—–] survives in any table cell (proves the cell-content
          owner sanitizes). save() itself auto-validates and would raise if a
          dash leaked, so a clean save is the first signal.
  GAP 3 — validate._check_dashes is a backstop. Craft a .docx DIRECTLY with
          python-docx (bypassing Report.table(), so no sanitize runs), inject
          an em-dash and an en-dash into table cells, and assert _check_dashes
          returns TIER_FAIL `dashes` violations located at the right cell.
          Also assert a clean table (and a clean prose-only doc) yields zero.
  GAP 4 — table() must NOT mangle dashes inside inline $...$ math in a cell.
          Cell sanitizing is delegated to _append_inline._prose(), which skips
          $...$ spans, so an em-dash inside `$a — b$` survives in the math
          <m:t> verbatim (identical to prose text()), while the prose segments
          of the same cell are still sanitized. Regression guard for an
          owner-level whole-string sanitize that corrupted in-cell math.

Exit codes:
    0  all tests passed
    1  at least one test failed
    2  environment missing (python-docx not importable)
"""
from __future__ import annotations

import re
import sys
import tempfile
import traceback
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))


def _require_env() -> None:
    try:
        import docx  # noqa: F401
    except ImportError as e:
        sys.stderr.write(
            f"SKIP-FATAL: python-docx not available ({e}). Build the skill venv\n"
            "  GOST_REPORT_EXTRAS=viz python3 scripts/ensure_env.py\n"
            "or run with ~/.local/share/agentpipe/gost-report/venv/bin/python.\n"
        )
        raise SystemExit(2)


_require_env()

from docx import Document  # noqa: E402

import gost_report as G  # noqa: E402
import validate as V  # noqa: E402

_PROHIBITED = re.compile(r"[—–]")
_MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# ----------------------------------------------------------------------------
# tiny test harness (same shape as test_viz.py)
# ----------------------------------------------------------------------------
_FAILURES: list[str] = []
_PASSED = 0


def check(name: str, fn) -> None:
    global _PASSED
    try:
        fn()
    except Exception:  # noqa: BLE001 — collect, don't abort the suite
        _FAILURES.append(name + "\n" + traceback.format_exc())
        sys.stderr.write(f"FAIL  {name}\n")
    else:
        _PASSED += 1
        print(f"ok    {name}")


def assert_eq(a, b, msg=""):
    assert a == b, f"{msg}: {a!r} != {b!r}"


_TMP = Path(tempfile.mkdtemp(prefix="sanitize_integ_"))


def _cell_texts(doc) -> list[str]:
    out: list[str] = []
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                out.append(c.text)
    return out


def _math_text_nodes(doc) -> list[str]:
    # Inline-math glyphs land in <m:t> nodes, NOT in cell.text. Sanitize must
    # never touch these (they carry LaTeX/math content verbatim).
    return [mt.text for mt in doc.element.body.iter(f"{{{_MATH_NS}}}t")]


# ----------------------------------------------------------------------------
# GAP 2 — Report.table() + text() sanitize on the real docx write path
# ----------------------------------------------------------------------------
def test_gap2_table_cells_sanitized_in_saved_docx():
    out = _TMP / "gap2.docx"
    r = G.Report(title_page=False)
    # text() with an em-dash (already covered by prose sanitize; included so the
    # save()'s auto-validation also sees a clean body).
    r.text("Измерения U — напряжение и I — ток.")
    # Header cells AND data cells carry em/en-dashes.
    r.table(
        [
            ["Параметр — обозн.", "Значение – диапазон"],   # header row
            ["U — напряжение", "5 — 10 В"],                  # data row
            ["I – ток", "0–3 А"],                            # en-dashes
        ],
        caption="Результаты",
    )
    # save() auto-validates; if any dash leaked into a cell it would raise
    # GostValidationError here. A clean return is the first proof.
    saved = r.save(out)
    assert saved.exists(), "save() did not write the .docx"

    doc = Document(str(saved))
    cells = _cell_texts(doc)
    leaks = [c for c in cells if _PROHIBITED.search(c)]
    assert not leaks, f"prohibited dash leaked into saved table cells: {leaks!r}"

    # Spot-check the canonical normalization of representative cells.
    assert "U, напряжение" in cells, f"expected normalized header cell; got {cells!r}"
    assert "5, 10 В" in cells, f"expected normalized data cell; got {cells!r}"
    # Single dash inside a range (no surrounding spaces) -> hyphen, not comma.
    assert "0-3 А" in cells, f"expected range dash -> hyphen; got {cells!r}"


def test_gap2_validator_clean_on_saved_docx():
    # Re-validate the GAP-2 output explicitly through _check_dashes: zero cell
    # violations (complements save()'s own auto-validation).
    out = _TMP / "gap2_revalidate.docx"
    r = G.Report(title_page=False)
    r.table([["A — B", "C – D"], ["e — f", "g – h"]], caption="Тест")
    saved = r.save(out)
    doc = Document(str(saved))
    viols = V._check_dashes(doc)
    dash_fails = [v for v in viols if v.code == "dashes"]
    assert not dash_fails, f"_check_dashes flagged sanitized cells: {dash_fails!r}"


# ----------------------------------------------------------------------------
# GAP 4 — table() must NOT mangle dashes inside inline $...$ math in a cell.
# Regression for a sanitize bug: an owner-level `value = _sanitize_prose(value)`
# ran on the WHOLE cell string BEFORE _append_inline split on the $...$ math
# delimiters, so an em-dash inside `$a — b$` was rewritten to a comma in the
# math (<m:t> ['a',',','b']) — breaking the "LaTeX/math untouched" contract and
# diverging from prose text(), where the same `$a — b$` keeps the dash. The fix
# defers all sanitizing to _append_inline._prose(), which skips $...$ spans.
# ----------------------------------------------------------------------------
def test_gap4_inline_math_dash_preserved_in_table_cell():
    out = _TMP / "gap4.docx"
    r = G.Report(title_page=False)
    r.table(
        [
            # prose with em-dash + trailing inline-math span carrying an em-dash
            ["Параметр — обозн.", r"Формула $a — b$"],
            # cell that is pure inline-math with an em-dash
            ["U — напряжение", r"$c — d$"],
        ],
        caption="Math-in-cell",
    )
    saved = r.save(out)  # auto-validates; would raise if a prose dash leaked
    assert saved.exists(), "save() did not write the .docx"

    doc = Document(str(saved))

    # 1) Prose segments are still sanitized (the regression must not over-correct
    #    by skipping prose). Including the prose PREFIX before an inline-math span.
    cells = _cell_texts(doc)
    leaks = [c for c in cells if _PROHIBITED.search(c)]
    assert not leaks, f"prose dash leaked into a cell: {leaks!r}"
    assert "Параметр, обозн." in cells, f"prose cell not sanitized: {cells!r}"
    assert "U, напряжение" in cells, f"prose cell not sanitized: {cells!r}"
    assert "Формула " in cells, f"prose prefix before math not sanitized: {cells!r}"

    # 2) The em-dash INSIDE $...$ survives verbatim in the math <m:t> nodes —
    #    identical to prose text() behavior, NOT mangled to a comma.
    math_nodes = _math_text_nodes(doc)
    assert "—" in math_nodes, (
        f"em-dash inside inline math was mangled (expected a verbatim '—' in "
        f"<m:t>); got {math_nodes!r}"
    )
    assert "," not in math_nodes, (
        f"sanitize leaked into inline math (comma found in <m:t>); "
        f"got {math_nodes!r}"
    )


def test_gap4_table_math_matches_prose_text_math():
    # Same inline-math span via table() vs via text(): the math <m:t> stream must
    # be byte-identical (proves the cell path no longer diverges from prose).
    expr = r"$x — y$"

    r_tbl = G.Report(title_page=False)
    r_tbl.table([[expr]], caption="Inline math", has_header=False)
    tbl_math = _math_text_nodes(Document(str(r_tbl.save(_TMP / "gap4_tbl.docx"))))

    r_txt = G.Report(title_page=False)
    r_txt.text(expr)
    txt_math = _math_text_nodes(Document(str(r_txt.save(_TMP / "gap4_txt.docx"))))

    assert_eq(tbl_math, txt_math,
              "table() inline-math <m:t> stream must equal text() inline-math")
    assert "—" in tbl_math, f"em-dash must survive in both paths; got {tbl_math!r}"


# ----------------------------------------------------------------------------
# GAP 3 — validator backstop: raw em-dash in a cell (bypassing table()) fails
# ----------------------------------------------------------------------------
def _build_dirty_docx_directly() -> "Document":
    # Build with python-docx directly so Report.table()'s sanitize never runs.
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    table.rows[1].cells[0].text = "U — напряжение"   # raw em-dash
    table.rows[1].cells[1].text = "5 – 10 В"          # raw en-dash
    return doc


def test_gap3_backstop_flags_raw_dash_in_cell():
    doc = _build_dirty_docx_directly()
    viols = V._check_dashes(doc)
    dash_fails = [v for v in viols if v.code == "dashes"]
    assert dash_fails, "backstop must flag a raw dash injected into a table cell"
    for v in dash_fails:
        assert_eq(v.tier, V.TIER_FAIL, "cell dash must be a hard FAIL")
        assert "таблица" in v.location, f"location must name the table: {v.location!r}"
    # The em-dash cell is row 2, col 1.
    locs = {v.location for v in dash_fails}
    assert "таблица 1, строка 2, столбец 1" in locs, \
        f"expected the em-dash cell location; got {locs!r}"
    assert "таблица 1, строка 2, столбец 2" in locs, \
        f"expected the en-dash cell location; got {locs!r}"


def test_gap3_backstop_clean_table_no_violation():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    table.rows[1].cells[0].text = "U, напряжение"   # already clean
    table.rows[1].cells[1].text = "5-10 В"
    viols = [v for v in V._check_dashes(doc) if v.code == "dashes"]
    assert not viols, f"clean table must yield zero cell violations: {viols!r}"


def test_gap3_paragraph_only_doc_clean():
    # Sanity: a prose-only doc with no tables and no dashes -> zero.
    doc = Document()
    doc.add_paragraph("Чистый текст без длинного и среднего тире.")
    viols = [v for v in V._check_dashes(doc) if v.code == "dashes"]
    assert not viols, f"clean prose doc must yield zero violations: {viols!r}"


def test_gap3_full_validate_docx_fails_on_dirty_cell():
    # End-to-end through the public validate_docx() entrypoint: a dirty cell
    # docx written to disk must produce a TIER_FAIL the same way save() would.
    dirty = _TMP / "gap3_dirty.docx"
    _build_dirty_docx_directly().save(str(dirty))
    viols = V.validate_docx(dirty)
    assert V.has_failures(viols), "validate_docx must report failures on dirty cell"
    assert any(v.code == "dashes" and v.tier == V.TIER_FAIL for v in viols), \
        f"expected a dashes TIER_FAIL; got {[(v.code, v.tier) for v in viols]!r}"


# ----------------------------------------------------------------------------
# runner
# ----------------------------------------------------------------------------
def main() -> int:
    # GAP 2
    check("gap2: table()+text() cells sanitized in saved .docx",
          test_gap2_table_cells_sanitized_in_saved_docx)
    check("gap2: _check_dashes clean on sanitized saved .docx",
          test_gap2_validator_clean_on_saved_docx)

    # GAP 4 — inline math in a cell is NOT mangled by sanitize
    check("gap4: inline $...$ dash preserved in a table cell",
          test_gap4_inline_math_dash_preserved_in_table_cell)
    check("gap4: table() inline-math == text() inline-math",
          test_gap4_table_math_matches_prose_text_math)

    # GAP 3
    check("gap3: backstop flags raw dash injected into a cell",
          test_gap3_backstop_flags_raw_dash_in_cell)
    check("gap3: backstop clean table -> zero cell violations",
          test_gap3_backstop_clean_table_no_violation)
    check("gap3: prose-only clean doc -> zero violations",
          test_gap3_paragraph_only_doc_clean)
    check("gap3: validate_docx() hard-fails a dirty-cell .docx",
          test_gap3_full_validate_docx_fails_on_dirty_cell)

    total = _PASSED + len(_FAILURES)
    if _FAILURES:
        sys.stderr.write("\n" + "=" * 70 + "\n")
        for f in _FAILURES:
            sys.stderr.write(f + "\n")
        sys.stderr.write(f"\n{len(_FAILURES)}/{total} test(s) FAILED.\n")
        return 1
    print(f"\nOK  all {total} integration tests passed")
    return 0


if __name__ == "__main__":
    sys.exit(main())
