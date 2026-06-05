"""
gost_report — генератор студенческих/научных работ по ГОСТ 7.32 для любого
российского вуза. Профили университетов задают поля, кегли заголовков,
название университета/факультета и прочие специфичные мелочи; чистый GOST 7.32
работает «из коробки», ИТМО — встроенный пресет.

Минимальный пример (ИТМО, дефолт):
    from gost_report import Report, TitleConfig

    r = Report(TitleConfig(
        work_type="Лабораторная работа",
        work_number="№1",
        topic="Основы работы в командной строке Unix",
        student_name="Фамилия И.О.",
        student_group="P3XXX",
        teacher_name="Фамилия И.О.",
        teacher_degree="к.т.н.",
        teacher_position="доцент",
        year="2026",
    ))
    r.toc()
    r.h1("Введение")
    r.text("Цель работы — изучить...")
    r.save("report.docx")

Любой другой вуз — передайте свой профиль:
    from gost_report import Report, TitleConfig, GOST_PROFILE, UniversityProfile

    SPBSU_PROFILE = UniversityProfile(
        university_short="«Санкт-Петербургский государственный университет»",
        faculty="Математико-механический факультет",
        toc_title="СОДЕРЖАНИЕ",
    )
    r = Report(TitleConfig(...), profile=SPBSU_PROFILE)

Зависимости: pip install python-docx
"""

from dataclasses import dataclass, field
from datetime import datetime, timezone
import os
from pathlib import Path
import re
import sys
from typing import Dict, List, Optional, Sequence, Union
import zipfile

from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt, Mm, Cm, Length, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _paths import paths, _caller_file, _find_root


# ============================================================
# Не варьируемые константы ГОСТ 7.32
# ============================================================

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(14)         # основной текст
FONT_SIZE_PAGE_NUMBER = Pt(11)  # номер страницы
FONT_SIZE_TITLE_LARGE = Pt(18)  # вид работы и тема на титульнике
LINE_SPACING_BODY = 1.5
FIRST_LINE_INDENT = Cm(1.25)    # абзацный отступ
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
PAGE_WIDTH = Mm(PAGE_WIDTH_MM)
PAGE_HEIGHT = Mm(PAGE_HEIGHT_MM)

# Line-spacing presets для тел разных типов абзацев.
LINE_SPACING_CODE = 1.0
LINE_SPACING_TABLE = 1.15

# Вертикальный отступ между блоками (figure caption, formula, list).
SPACE_BLOCK = Pt(6)

# Титульник: пустые абзацы-разделители и сдвиг правого блока.
_TITLE_TOP_GAP = 6                    # пустых абзацев от шапки до вида работы
_TITLE_MID_GAP = 4                    # пустых абзацев перед блоком студент/преподаватель
_RIGHT_BLOCK_INDENT = Cm(9)           # сдвиг правого блока от левого края
_RIGHT_BLOCK_REST_INDENT = Cm(2.2)    # доп. сдвиг continuation-имён (≈ ширина "Выполнили: ")


# ============================================================
# Профиль университета — всё, что варьируется между вузами
# ============================================================

@dataclass
class UniversityProfile:
    """Параметры оформления, специфичные для вуза.

    Дефолты соответствуют ГОСТ 7.32-2017 в наиболее распространённом
    толковании. Конкретные вузы (ИТМО, МГУ, СПбГУ и т.д.) переопределяют
    нужные поля. Создавайте свой профиль для своего вуза или принимайте
    PR-ы с новыми константами в этом модуле.
    """
    # Поля титульного листа (мм)
    title_margin_left: float = 30
    title_margin_right: float = 15
    title_margin_top: float = 20
    title_margin_bottom: float = 20
    # Поля основного текста (мм)
    body_margin_left: float = 30
    body_margin_right: float = 15
    body_margin_top: float = 20
    body_margin_bottom: float = 20
    # Кегли заголовков (пт)
    heading_size_h1: int = 16
    heading_size_h2: int = 14
    heading_size_h3: int = 14
    # Выравнивание заголовков: "center" | "left"
    heading_align_h1: str = "center"
    heading_align_h2: str = "left"
    heading_align_h3: str = "left"
    # Поведение h1
    h1_uppercase: bool = True
    h1_new_page: bool = True
    # Заголовок оглавления — ГОСТ 7.32-2017 предписывает «СОДЕРЖАНИЕ»,
    # но многие вузы (ИТМО) используют «ОГЛАВЛЕНИЕ».
    toc_title: str = "СОДЕРЖАНИЕ"
    # Поля титульника, относящиеся к вузу — fallback'и для TitleConfig
    ministry: str = "Министерство науки и высшего образования Российской Федерации"
    university_full: str = ""
    university_short: str = ""
    faculty: str = ""
    city: str = ""


# ---- Встроенные пресеты ----

GOST_PROFILE = UniversityProfile()
"""Чистый ГОСТ 7.32-2017. Поля университета не заданы — должны
прийти из TitleConfig либо из вашего собственного профиля."""

ITMO_PROFILE = UniversityProfile(
    title_margin_left=30,
    title_margin_right=20,
    title_margin_top=10,
    title_margin_bottom=10,
    body_margin_left=30,
    body_margin_right=10,
    body_margin_top=20,
    body_margin_bottom=20,
    heading_size_h1=16,
    heading_size_h2=15,
    heading_size_h3=14,
    heading_align_h1="center",
    heading_align_h2="center",
    heading_align_h3="center",
    toc_title="ОГЛАВЛЕНИЕ",
    university_full=(
        "федеральное государственное автономное "
        "образовательное учреждение высшего образования"
    ),
    university_short="«Национальный исследовательский университет ИТМО»",
    faculty="Факультет программной инженерии и компьютерной техники",
    city="Санкт-Петербург",
)
"""Университет ИТМО, ФПИиКТ. Используется как дефолт, чтобы не ломать
существующие лабораторные."""

DEFAULT_PROFILE = ITMO_PROFILE


# Sentinel for "do not set this attribute — inherit from the paragraph style".
# Used by _make_paragraph callers (notably list paragraphs, which must keep
# w:before unset so the value comes from the List Number / List Bullet style).
_INHERIT = object()


# ============================================================
# Стилизованный run внутри одного абзаца
# ============================================================

@dataclass(frozen=True)
class StyledRun:
    """Кусочек текста с собственным форматированием, для абзацев из
    нескольких run'ов (например "Группа: P3XXX" где "Группа:" подчёркнут
    а номер — обычный). Используется на титульнике.
    """
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    # docx.shared.Length (либо None → FONT_SIZE_BODY). Pt/Cm/Mm возвращают
    # Length-объекты; используем Optional[Length] вместо object для ясности.
    size: Optional["Length"] = None


# ============================================================
# Ошибка валидации
# ============================================================

class GostValidationError(RuntimeError):
    """Поднимается из Report.save() когда сгенерированный .docx не проходит
    автоматическую проверку ГОСТ.

    Текст исключения — единственная prose-инструкция про валидацию во всём
    скилле. Она появляется в shell-выводе только когда модель уже накосячила.
    Поэтому фразировка имеет значение: явно сказано «перегенерируй», «не
    подавляй», «не используй python-docx напрямую» — это именно те ошибочные
    реакции, которые надо превентить.
    """


# ============================================================
# Конфигурация титульного листа (контент, не оформление)
# ============================================================

@dataclass
class TitleConfig:
    """Содержимое титульного листа.

    Обязательные по смыслу: work_type, topic, student_name (или student_names),
    student_group, year. Любое из них можно опустить здесь и задать через
    env-переменные (`~/.config/gost-report/config` или `os.environ`), см.
    `_load_gost_env`. Env побеждает build.py, если задан и не пуст.

    Поля университета (university_*, faculty, city, ministry) можно оставить
    пустыми — тогда они подтянутся из активного UniversityProfile.
    """
    # --- Обязательное по смыслу (env может заполнить) ---
    work_type: str = ""
    topic: str = ""
    student_name: str = ""
    student_group: str = ""
    year: str = ""

    # --- Команда (если задан student_names — рендерится список) ---
    student_names: List[str] = field(default_factory=list)
    student_label: str = ""                # пусто = авто: "Выполнили"/"Выполнил"

    # --- Опциональное ---
    work_number: str = ""
    variant: str = ""
    teacher_name: str = ""
    teacher_label: str = "Проверил"        # "Руководитель" для ВКР/курсовых
    teacher_degree: str = ""
    teacher_position: str = ""

    # --- Переопределение полей профиля (пусто = взять из профиля) ---
    city: str = ""
    ministry: str = ""
    university_full: str = ""
    university_short: str = ""
    faculty: str = ""


# ============================================================
# Env-конфиг: ~/.config/gost-report/config + os.environ
# ============================================================

# Ключи env, которые накладываются на TitleConfig. Значение — имя атрибута
# TitleConfig. Префикс "GOST_REPORT_" + UPPER(имя атрибута) для каждого.
_ENV_PREFIX = "GOST_REPORT_"
_ENV_STR_FIELDS = (
    "work_type", "topic", "student_name", "student_group", "year",
    "student_label",
    "work_number", "variant",
    "teacher_name", "teacher_label", "teacher_degree", "teacher_position",
    "city", "ministry", "university_full", "university_short", "faculty",
)


def _config_file_path() -> Path:
    """Стандартный путь конфига: $XDG_CONFIG_HOME/gost-report/config
    либо ~/.config/gost-report/config. Можно переопределить через
    GOST_REPORT_CONFIG."""
    override = os.environ.get(f"{_ENV_PREFIX}CONFIG", "").strip()
    if override:
        return Path(override).expanduser()
    base = os.environ.get("XDG_CONFIG_HOME", "").strip()
    if base:
        return Path(base).expanduser() / "gost-report" / "config"
    return Path.home() / ".config" / "gost-report" / "config"


_CONFIG_LINE_RE = re.compile(r'^\s*(?:export\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*=\s*(.*?)\s*$')


def _parse_config_file(path: Path) -> Dict[str, str]:
    """Простой парсер .env-формата: KEY=VALUE построчно. Поддерживает
    `export KEY=...`, кавычки (' или "), комментарии (#) и пустые строки.
    Никаких подстановок переменных или escape-последовательностей."""
    if not path.is_file():
        return {}
    out: Dict[str, str] = {}
    try:
        text = path.read_text(encoding="utf-8")
    except OSError:
        return {}
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        m = _CONFIG_LINE_RE.match(raw_line)
        if not m:
            continue
        key, value = m.group(1), m.group(2)
        if len(value) >= 2 and value[0] == value[-1] and value[0] in ('"', "'"):
            value = value[1:-1]
        out[key] = value
    return out


_SUBJECT_CONFIG_NAME = ".gost-report.env"


def _find_subject_config() -> Optional[Path]:
    """Subject-уровневый конфиг: обход вверх от build.py до project root,
    ищем `.gost-report.env`. Возвращает ближайший. Покрывает и
    отдельные репы по предмету, и монорепо с несколькими предметами:
    конфиг в папке `физика/` побеждает конфиг в корне монорепо."""
    start = _caller_file() or Path.cwd()
    if start.is_file():
        start = start.parent
    start = start.resolve()
    root = _find_root(start)
    upper = root.resolve() if root else start

    cur = start
    while True:
        candidate = cur / _SUBJECT_CONFIG_NAME
        if candidate.is_file():
            return candidate
        if cur == upper or cur.parent == cur:
            break
        cur = cur.parent
    return None


def _load_gost_env() -> Dict[str, str]:
    """Сливает три слоя env: global config → subject config → os.environ.
    Каждый последующий перекрывает предыдущий. Пустые/отсутствующие
    значения не затирают то, что выставлено выше по приоритету."""
    merged = _parse_config_file(_config_file_path())
    subject = _find_subject_config()
    if subject is not None:
        merged.update(_parse_config_file(subject))
    for key, val in os.environ.items():
        if key.startswith(_ENV_PREFIX) and val:
            merged[key] = val
    return merged


# Маркер заглушек из install-template config.env.example. Значение, содержащее
# этот токен, трактуется как НЕзаполненное везде — чтобы первый отчёт не ушёл с
# «ЗАПОЛНИ ФИО» на титуле, а валидатор указал, где править.
_PLACEHOLDER_TOKEN = "ЗАПОЛНИ"


def _is_placeholder(value: str) -> bool:
    return _PLACEHOLDER_TOKEN in (value or "").upper()


def _apply_env_overrides(cfg: TitleConfig) -> TitleConfig:
    """Накладывает env поверх TitleConfig. Env побеждает build.py, если
    непустое и не заглушка. Пустое/отсутствующее/заглушка-env — оставляет
    значение из build.py.

    Особый случай: GOST_REPORT_STUDENT_NAMES — список через запятую,
    перекрывает только cfg.student_names (не student_name).
    """
    env = _load_gost_env()
    for field_name in _ENV_STR_FIELDS:
        env_key = f"{_ENV_PREFIX}{field_name.upper()}"
        env_val = env.get(env_key, "")
        if env_val and not _is_placeholder(env_val):
            setattr(cfg, field_name, env_val)
    names_raw = env.get(f"{_ENV_PREFIX}STUDENT_NAMES", "")
    if names_raw and not _is_placeholder(names_raw):
        cfg.student_names = [n.strip() for n in names_raw.split(",") if n.strip()]
    return cfg


_TRUTHY_ENV = {"1", "true", "yes", "on", "да"}


def _env_flag(name: str) -> bool:
    """Булев env-флаг из слитого gost-env (config-файлы + os.environ).
    Префикс GOST_REPORT_ добавляется автоматически. Truthy: 1/true/yes/on/да."""
    return _load_gost_env().get(
        f"{_ENV_PREFIX}{name}", "").strip().lower() in _TRUTHY_ENV


def _resolve_student_label(cfg: TitleConfig) -> str:
    """Авто-выбор: команда — "Выполнили", иначе — "Выполнил".
    Явно заданное cfg.student_label всегда побеждает."""
    if cfg.student_label:
        return cfg.student_label
    if cfg.student_names:
        return "Выполнили"
    return "Выполнил"


# ============================================================
# Низкоуровневые помощники работы с XML python-docx
# ============================================================

_RFONTS_ATTRS = ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia")


def _force_run_fonts(rpr, font_name: str) -> None:
    """Принудительно ставит ``font_name`` во все <w:rFonts> слоты (ascii,
    hAnsi, cs, eastAsia). Используется в стилях и run-level форматировании
    чтобы Word гарантированно отрисовал нужный шрифт для всех кодовых
    диапазонов (латиница, кириллица, восточно-азиатские, complex)."""
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in _RFONTS_ATTRS:
        rfonts.set(qn(attr), font_name)


def _set_run_font(run, *, size=FONT_SIZE_BODY, bold=False, italic=False,
                  underline=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    _force_run_fonts(run._element.get_or_add_rPr(), FONT_NAME)


_THUMBNAIL_PART_NAME = "docProps/thumbnail.jpeg"
_THUMBNAIL_REL_RE = re.compile(
    rb'<Relationship[^>]*Target="docProps/thumbnail\.jpeg"[^>]*/>'
)


# Фиксированная дата для детерминированного docProps/core.xml. python-docx
# проставляет dcterms:created/modified текущим временем при save(), из-за чего
# один и тот же документ, собранный в двух процессах, даёт разные байты core.xml
# (остальные part-ы уже детерминированы). Пиним на константу; уважаем
# SOURCE_DATE_EPOCH (конвенция reproducible-builds), если задан.
_FALLBACK_TIMESTAMP = datetime(2000, 1, 1, 0, 0, 0)  # naive UTC


def _deterministic_timestamp() -> datetime:
    """Naive-UTC datetime для core.xml: SOURCE_DATE_EPOCH если задан, иначе константа."""
    raw = os.environ.get("SOURCE_DATE_EPOCH")
    if raw:
        try:
            return datetime.fromtimestamp(
                int(raw), tz=timezone.utc
            ).replace(tzinfo=None)
        except (ValueError, OverflowError, OSError):
            pass
    return _FALLBACK_TIMESTAMP


def _strip_python_docx_fingerprints(doc: _Document) -> None:
    """Стирает дефолтные строки шаблона python-docx из core_properties:
    «python-docx» в dc:creator, «generated by python-docx» в dc:description,
    фоссилизированный таймстемп 2013-12-23 в dcterms:created/modified.
    Без этого любой, кто откроет «Свойства документа» в Word или распакует
    .docx как zip, сразу видит, что файл сгенерирован библиотекой.

    Таймстемпы пинятся на фиксированную дату (см. ``_deterministic_timestamp``),
    а cp:revision/lastModifiedBy нормализуются, чтобы core.xml был байт-в-байт
    воспроизводимым между процессами.
    """
    cp = doc.core_properties
    cp.author = ""
    cp.comments = ""
    cp.last_modified_by = ""
    cp.revision = 1
    pinned = _deterministic_timestamp()
    cp.created = pinned
    cp.modified = pinned


def _strip_thumbnail_part(docx_path: Path) -> None:
    """Удаляет дефолтный thumbnail python-docx из сохранённого .docx.

    Эта jpeg-картинка (8.3 КБ, 395x512) одинакова в каждом python-docx файле,
    однозначно идентифицирует библиотеку и не появляется в обычных Word-файлах,
    сохранённых без галки «Save Thumbnail». Удаляем сам part и Relationship
    на него в `_rels/.rels`; запись `<Default Extension="jpeg"/>` в
    `[Content_Types].xml` оставляем — без привязанного файла она безвредна
    и встречается в Word-файлах с jpeg-картинками внутри (figure-ы).
    """
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            names = set(zin.namelist())
            if _THUMBNAIL_PART_NAME not in names:
                return
            entries = [
                (item, zin.read(item.filename))
                for item in zin.infolist()
                if item.filename != _THUMBNAIL_PART_NAME
            ]
    except (zipfile.BadZipFile, OSError):
        return
    tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")
    try:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item, data in entries:
                if item.filename == "_rels/.rels":
                    data = _THUMBNAIL_REL_RE.sub(b"", data)
                zout.writestr(item, data)
        tmp.replace(docx_path)
    except Exception:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass
        raise


def _ensure_normal_style(doc: _Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE_BODY
    _force_run_fonts(normal.element.get_or_add_rPr(), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING_BODY
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def _append_word_field(run, instr_text: str, *,
                       placeholder: Optional[str] = None,
                       preserve_space: bool = False) -> None:
    """Вставляет Word-поле в ``run``: ``<w:fldChar begin>`` + ``<w:instrText>``
    + (опционально) ``<w:fldChar separate>`` + placeholder-run + ``<w:fldChar end>``.

    Используется для PAGE (footer) и TOC (заголовок оглавления). Если задан
    ``placeholder`` — добавляется separate-маркер и placeholder-run
    с текстом, который Word/Pages/LibreOffice заменят при обновлении полей.
    ``preserve_space=True`` ставит ``xml:space="preserve"`` на ``instrText``
    (нужно для TOC, где переключатели разделены пробелами).
    """
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    if preserve_space:
        instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text

    run._element.append(fld_begin)
    run._element.append(instr)

    if placeholder is not None:
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        placeholder_run = OxmlElement("w:r")
        placeholder_text = OxmlElement("w:t")
        placeholder_text.text = placeholder
        placeholder_run.append(placeholder_text)
        run._element.append(fld_separate)
        run._element.append(placeholder_run)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)


def _add_page_number_to_footer(section, *, hide_on_first=True):
    if hide_on_first:
        section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False

    for p in list(footer.paragraphs):
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING_CODE

    run = p.add_run()
    _set_run_font(run, size=FONT_SIZE_PAGE_NUMBER)
    _append_word_field(run, "PAGE")

    if hide_on_first:
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        for p in list(first_footer.paragraphs):
            p.clear()


def _set_section_a4(section):
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT


def _apply_margins(section, left, right, top, bottom):
    section.left_margin = Mm(left)
    section.right_margin = Mm(right)
    section.top_margin = Mm(top)
    section.bottom_margin = Mm(bottom)


def _align_const(name: str):
    return WD_ALIGN_PARAGRAPH.CENTER if name == "center" else WD_ALIGN_PARAGRAPH.LEFT


def _configure_heading_style(doc: _Document, level: int, size_pt: int,
                             align: str):
    style = doc.styles[f"Heading {level}"]
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    _force_run_fonts(style.element.get_or_add_rPr(), FONT_NAME)

    pf = style.paragraph_format
    pf.alignment = _align_const(align)
    pf.line_spacing = LINE_SPACING_BODY
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True


# ============================================================
# Санитайзер пользовательской прозы
# ============================================================
#
# Любой текст, который агент передаёт в text/task/numbered/bullet/h1-h3 и
# в аргумент caption у figure/table, проходит через _sanitize_prose. Длинное
# и среднее тире — типичный AI-маркер и прямо запрещены writing-style блоком
# в SKILL.md, но модели всё равно их иногда вставляют. Подстраховываемся
# на уровне библиотеки: " — " / " – " между словами становится ", ",
# одиночные тире (например в диапазонах "1—5") — обычным дефисом.
#
# Авто-префиксы "Рисунок N — Описание" и "Таблица N — Описание" формирует
# сама библиотека — там тире обязательно по ГОСТ и не трогается.

_PROSE_DASH_BETWEEN_WORDS = re.compile(r"\s+[—–]\s+")
_PROSE_DASH_REMAINING = re.compile(r"[—–]")


def _sanitize_prose(text):
    if not text:
        return text
    text = _PROSE_DASH_BETWEEN_WORDS.sub(", ", text)
    text = _PROSE_DASH_REMAINING.sub("-", text)
    return text


# LaTeX → OMML вынесено в модуль gost_report_math (см. research/19 §5).
# Формулы доступны через r.formula(...) (BC-shim → r.math.formula).


# ============================================================
# Класс отчёта — высокоуровневый API
# ============================================================

class Report:
    """Главный класс. Собирает документ из вызовов h1/text/code/figure/...

    Создание:
        r = Report(TitleConfig(...))                    # ITMO_PROFILE по умолчанию
        r = Report(TitleConfig(...), profile=GOST_PROFILE)
        r = Report(TitleConfig(...), profile=UniversityProfile(...))
        r = Report(title_page=False)                    # без титульника (TitleConfig опционален)

    Без титульного листа: `Report(title_page=False)` (или env
    `GOST_REPORT_NO_TITLE_PAGE=1`) — документ начинается сразу с основного
    текста, поля основного текста, нумерация с первой страницы. Поля
    TitleConfig в этом режиме не обязательны (валидация work_type/topic/ФИО/
    группы/года пропускается).

    Контент: см. методы toc/h1/h2/h3/text/task/code/figure/table/numbered/bullet.
    Сохранение: r.save("report.docx").
    """

    def __init__(self, title: Optional[TitleConfig] = None,
                 profile: UniversityProfile = DEFAULT_PROFILE,
                 *,
                 title_page: bool = True,
                 project_root: Optional[Union[str, Path]] = None):
        self._doc = Document()
        # Титульник можно отключить: параметр title_page=False или env
        # GOST_REPORT_NO_TITLE_PAGE=1. Env только отключает (не включает),
        # чтобы автоматизация могла гасить титул, не трогая build.py.
        self._has_title_page = title_page and not _env_flag("NO_TITLE_PAGE")
        # Env побеждает build.py: накладываем перед валидацией обязательных
        # полей, чтобы можно было опускать FIO/group/year в TitleConfig.
        self._title = _apply_env_overrides(
            title if title is not None else TitleConfig())
        # Обязательные поля титула нужны только когда титул строится.
        if self._has_title_page:
            self._check_required(self._title)
        self._profile = profile
        self._figure_counter = 0
        self._table_counter = 0
        self._formula_counter = 0
        self._just_broke_page = False
        self._next_num_id = 100

        # Подключаемые модули (r.plot, r.diagram, …). discover() лёгкий — тянет
        # лишь классы модулей (core_api), тяжёлые пакеты грузятся при первом
        # использовании namespace через __getattr__ → attach(). См. registry.py.
        try:
            from registry import discover as _discover_modules
            self._modules = _discover_modules()
        except Exception:
            self._modules = {}
        self._attached = {}
        self._tmpdir = None

        # Авто-резолв путей: если project_root не задан — paths() обходит
        # стек вверх до user-скрипта, оттуда ищет маркер проекта (.git,
        # Makefile, pyproject.toml, .claude). Используется figure() для
        # резолва относительных имён файлов и save() для дефолтного пути.
        if project_root is None:
            self._paths = paths()
        else:
            self._paths = paths(Path(project_root))

        # 1. Базовые стили
        _ensure_normal_style(self._doc)
        _configure_heading_style(self._doc, 1,
                                 profile.heading_size_h1,
                                 profile.heading_align_h1)
        _configure_heading_style(self._doc, 2,
                                 profile.heading_size_h2,
                                 profile.heading_align_h2)
        _configure_heading_style(self._doc, 3,
                                 profile.heading_size_h3,
                                 profile.heading_align_h3)

        if self._has_title_page:
            # 2. Секция титульника (первая)
            section = self._doc.sections[0]
            _set_section_a4(section)
            _apply_margins(section,
                           profile.title_margin_left, profile.title_margin_right,
                           profile.title_margin_top, profile.title_margin_bottom)
            _add_page_number_to_footer(section, hide_on_first=True)
            self._build_title_page()

            # 3. Новая секция для основного текста
            body_section = self._doc.add_section(WD_SECTION.NEW_PAGE)
            _set_section_a4(body_section)
            _apply_margins(body_section,
                           profile.body_margin_left, profile.body_margin_right,
                           profile.body_margin_top, profile.body_margin_bottom)
            _add_page_number_to_footer(body_section, hide_on_first=False)
        else:
            # Без титульника: единственная секция сразу с полями основного
            # текста; нумерация страниц видна с первой страницы.
            section = self._doc.sections[0]
            _set_section_a4(section)
            _apply_margins(section,
                           profile.body_margin_left, profile.body_margin_right,
                           profile.body_margin_top, profile.body_margin_bottom)
            _add_page_number_to_footer(section, hide_on_first=False)
        self._just_broke_page = True

    # --------------------------------------------------------
    # Внутреннее: построение титульника
    # --------------------------------------------------------

    @staticmethod
    def _check_required(cfg: TitleConfig) -> None:
        """После применения env-override проверяет, что обязательные по
        смыслу поля заданы. Сообщение указывает на оба источника — build.py
        и env, чтобы пользователь сразу понимал где править."""
        def _missing(value: str) -> bool:
            return not value or _is_placeholder(value)

        problems = []
        if _missing(cfg.work_type):
            problems.append("work_type (GOST_REPORT_WORK_TYPE)")
        if _missing(cfg.topic):
            problems.append("topic (GOST_REPORT_TOPIC)")
        if _missing(cfg.student_name) and not cfg.student_names:
            problems.append(
                "student_name или student_names "
                "(GOST_REPORT_STUDENT_NAME / GOST_REPORT_STUDENT_NAMES)"
            )
        if _missing(cfg.student_group):
            problems.append("student_group (GOST_REPORT_STUDENT_GROUP)")
        if _missing(cfg.year):
            problems.append("year (GOST_REPORT_YEAR)")
        if problems:
            raise ValueError(
                "TitleConfig: не заданы обязательные поля: "
                + ", ".join(problems)
                + f". Заполни их в {_config_file_path()} (env-формат), "
                  "в <курс>/.gost-report.env, через переменные окружения "
                  "или прямо в TitleConfig(...). "
                  "Значения-заглушки «ЗАПОЛНИ…» считаются незаполненными."
            )

    def _resolve(self, attr: str) -> str:
        """TitleConfig.<attr> если задано, иначе UniversityProfile.<attr>."""
        value = getattr(self._title, attr, "") or ""
        if value:
            return value
        return getattr(self._profile, attr, "") or ""

    def _printable_cm(self) -> float:
        """Печатная ширина A4 минус поля активного профиля, в см.
        Используется figure() для clamp ширины картинки и formula() для
        размещения номера формулы на правом крае."""
        return (PAGE_WIDTH_MM
                - self._profile.body_margin_left
                - self._profile.body_margin_right) / 10.0

    # --------------------------------------------------------
    # Подключаемые модули: lazy-attach + CoreServices-фасад
    # --------------------------------------------------------
    def __getattr__(self, name: str):
        """Ленивое присоединение модуля при первом обращении к r.<namespace>.
        Срабатывает только если обычный lookup не нашёл атрибут. Ошибка
        зависимостей (ActionableImportError из check_available) пробрасывается
        наружу как есть, не превращается в AttributeError."""
        if name.startswith("_"):
            raise AttributeError(name)
        modules = self.__dict__.get("_modules") or {}
        if name in modules:
            attached = self.__dict__.setdefault("_attached", {})
            if name not in attached:
                attached[name] = modules[name].attach(self)
            return attached[name]
        raise AttributeError(name)

    def printable_cm(self) -> float:
        """Публичный фасад _printable_cm() для модулей."""
        return self._printable_cm()

    def sanitize(self, text: str) -> str:
        """Санитайз прозы (подписи/контент) — фасад _sanitize_prose."""
        return _sanitize_prose(text)

    def next_formula_number(self) -> int:
        """Следующий номер формулы (для math-модуля)."""
        self._formula_counter += 1
        return self._formula_counter

    def make_paragraph(self, **kwargs):
        """Пустой параграф с форматированием — фасад _make_paragraph (для
        модулей, строящих низкоуровневую раскладку, напр. формулы)."""
        return self._make_paragraph(**kwargs)

    def add_text_paragraph(self, text: str = "", **kwargs):
        """Текстовый параграф — фасад _add_paragraph."""
        return self._add_paragraph(text, **kwargs)

    def set_run_font(self, run, **kwargs) -> None:
        """Шрифт для run — фасад _set_run_font."""
        _set_run_font(run, **kwargs)

    @property
    def space_block(self):
        """Межблочный интервал SPACE_BLOCK (для модулей)."""
        return SPACE_BLOCK

    def embed_figure(self, image, caption: str, *,
                     width_cm: Optional[float] = None) -> int:
        """Единая точка нумерации+подписи рисунков для модулей. Принимает путь
        ИЛИ Figure-объект (chart/diagram). Возвращает номер рисунка."""
        self.figure(image, caption, width_cm=width_cm)
        return self._figure_counter

    def embed_table(self, rows: Sequence[Sequence[str]], caption: str = "",
                    *, has_header: bool = True) -> int:
        """Единая точка нумерации+подписи таблиц для модулей."""
        self.table(rows, caption=caption, has_header=has_header)
        return self._table_counter

    @property
    def tmp_dir(self) -> Path:
        """Временная директория для PNG, сгенерированных модулями. Чистится в
        save()/teardown. Путь в .docx не утекает — картинки встраиваются байтами."""
        if self.__dict__.get("_tmpdir") is None:
            import tempfile
            self._tmpdir = Path(tempfile.mkdtemp(prefix="gost-report-"))
        return self._tmpdir

    @property
    def paths(self):
        """ProjectPaths активного проекта (root/docs/figures/...)."""
        return self._paths

    def _cleanup(self) -> None:
        """Idempotent teardown подключённых модулей + удаление tmp-директории."""
        modules = self.__dict__.get("_modules") or {}
        for ns in list(self.__dict__.get("_attached") or {}):
            try:
                modules[ns].teardown()
            except Exception:
                pass
        tmp = self.__dict__.get("_tmpdir")
        if tmp:
            import shutil
            shutil.rmtree(tmp, ignore_errors=True)
            self._tmpdir = None

    def __del__(self):
        try:
            self._cleanup()
        except Exception:
            pass

    def _make_paragraph(self, *, align=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=LINE_SPACING_BODY,
                        space_before=Pt(0), space_after=Pt(0),
                        left_indent=None, first_line_indent=None):
        """Создаёт пустой параграф с заданным форматированием.

        Единая точка входа для всех «телом-уровневых» методов
        (text/task/figure/formula/table/caption/list). Возвращает Paragraph
        без runs — caller добавляет содержимое (текст, OMML, picture, и т.п.).

        Любой параметр (``align``, ``line_spacing``, ``space_before``,
        ``space_after``, ``left_indent``, ``first_line_indent``) можно явно
        обнулить через ``_INHERIT`` чтобы оставить значение из стиля абзаца
        (нужно для list-параграфов: их ``w:before`` берётся из стиля
        ``List Number`` / ``List Bullet``).
        """
        p = self._doc.add_paragraph()
        if align is not _INHERIT:
            p.alignment = align
        pf = p.paragraph_format
        if line_spacing is not _INHERIT:
            pf.line_spacing = line_spacing
        if space_before is not _INHERIT:
            pf.space_before = space_before
        if space_after is not _INHERIT:
            pf.space_after = space_after
        if left_indent is not None and left_indent is not _INHERIT:
            pf.left_indent = left_indent
        if first_line_indent is not None and first_line_indent is not _INHERIT:
            pf.first_line_indent = first_line_indent
        return p

    def _add_paragraph(self, text="", *, align=WD_ALIGN_PARAGRAPH.CENTER,
                       size=FONT_SIZE_BODY, bold=False, italic=False,
                       underline=False, left_indent=None,
                       first_line_indent=None,
                       line_spacing=LINE_SPACING_BODY,
                       space_before=Pt(0), space_after=Pt(0)):
        p = self._make_paragraph(align=align, line_spacing=line_spacing,
                                 space_before=space_before,
                                 space_after=space_after,
                                 left_indent=left_indent,
                                 first_line_indent=first_line_indent)
        if text:
            run = p.add_run(text)
            _set_run_font(run, size=size, bold=bold, italic=italic,
                          underline=underline)
        return p

    def _add_runs_paragraph(self, runs, *, align=WD_ALIGN_PARAGRAPH.CENTER,
                            left_indent=None):
        p = self._make_paragraph(align=align, left_indent=left_indent)
        for r in runs:
            run = p.add_run(r.text)
            # Explicit None-check, не `or`-fallback: иначе Pt(0) (бывает
            # falsy) был бы перетёрт дефолтом.
            run_size = r.size if r.size is not None else FONT_SIZE_BODY
            _set_run_font(
                run,
                size=run_size,
                bold=r.bold,
                italic=r.italic,
                underline=r.underline,
            )
        return p

    def _add_bottom_anchored_block(self, lines):
        """Borderless single-column таблица, плавающе прижатая к нижнему
        полю текущей страницы. Каждый элемент ``lines`` — отдельная строка
        (центр, TNR 14). Используется для города/года на титульнике, чтобы
        они оставались на 1-й странице независимо от длины шапки/темы/ФИО.

        Реализация — через OOXML ``w:tblpPr`` (плавающее позиционирование):
        ``vertAnchor=margin`` + ``tblpYSpec=bottom`` прижимает низ таблицы к
        нижнему полю; ``horzAnchor=margin`` + ``tblpXSpec=center`` центрирует
        по горизонтали внутри полей. python-docx высокоуровневого API для
        этого не даёт, поэтому правим XML напрямую.
        """
        table = self._doc.add_table(rows=len(lines), cols=1)
        table.autofit = True

        tbl = table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblpPr = OxmlElement("w:tblpPr")
        tblpPr.set(qn("w:vertAnchor"), "margin")
        tblpPr.set(qn("w:horzAnchor"), "margin")
        tblpPr.set(qn("w:tblpYSpec"), "bottom")
        tblpPr.set(qn("w:tblpXSpec"), "center")
        tblpPr.set(qn("w:leftFromText"), "0")
        tblpPr.set(qn("w:rightFromText"), "0")
        tblpPr.set(qn("w:topFromText"), "0")
        tblpPr.set(qn("w:bottomFromText"), "0")
        tblPr.append(tblpPr)

        tblOverlap = OxmlElement("w:tblOverlap")
        tblOverlap.set(qn("w:val"), "never")
        tblPr.append(tblOverlap)

        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        for row, text in zip(table.rows, lines):
            cell = row.cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING_BODY
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            run = p.add_run(text)
            _set_run_font(run, size=FONT_SIZE_BODY)

        return table

    def _build_title_page(self):
        self._title_header()
        self._add_blank_paragraphs(_TITLE_TOP_GAP)
        self._title_work_block()
        self._add_blank_paragraphs(_TITLE_MID_GAP)
        self._title_student_block()
        self._title_teacher_block()
        self._title_footer()

    def _add_blank_paragraphs(self, count: int) -> None:
        for _ in range(count):
            self._add_paragraph()

    def _title_header(self) -> None:
        """Министерство → университет → факультет."""
        ministry = self._resolve("ministry")
        if ministry:
            self._add_paragraph(ministry)
        university_full = self._resolve("university_full")
        if university_full:
            self._add_paragraph(university_full)
        university_short = self._resolve("university_short")
        if university_short:
            self._add_paragraph(university_short)
        faculty = self._resolve("faculty")
        if faculty:
            self._add_paragraph()
            self._add_paragraph(faculty, italic=True)

    def _title_work_block(self) -> None:
        """Вид работы (+ номер) → тема → вариант."""
        cfg = self._title
        work_type_text = cfg.work_type
        if cfg.work_number:
            work_type_text = f"{work_type_text} {cfg.work_number}".strip()
        self._add_paragraph(work_type_text, size=FONT_SIZE_TITLE_LARGE,
                            bold=True)
        if cfg.topic:
            self._add_paragraph(cfg.topic, size=FONT_SIZE_TITLE_LARGE)
        if cfg.variant:
            self._add_paragraph()
            self._add_paragraph(f"Вариант №{cfg.variant}",
                                size=FONT_SIZE_TITLE_LARGE)

    def _title_student_block(self) -> None:
        """Группа → "Выполнил(и): ФИО" с правым сдвигом."""
        cfg = self._title
        self._add_runs_paragraph(
            [StyledRun(f"Группа: {cfg.student_group}", underline=True)],
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=_RIGHT_BLOCK_INDENT,
        )
        student_label = _resolve_student_label(cfg)
        names = cfg.student_names or ([cfg.student_name] if cfg.student_name else [])
        first, *rest = names
        self._add_runs_paragraph(
            [
                StyledRun(student_label, underline=True),
                StyledRun(f": {first}"),
            ],
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=_RIGHT_BLOCK_INDENT,
        )
        # Дополнительные участники: визуально выравниваются под первое имя
        # через немного больший left_indent (≈ ширина "Выполнили: " в TNR 14pt).
        rest_indent = _RIGHT_BLOCK_INDENT + _RIGHT_BLOCK_REST_INDENT
        for extra_name in rest:
            self._add_runs_paragraph(
                [StyledRun(extra_name)],
                align=WD_ALIGN_PARAGRAPH.LEFT,
                left_indent=rest_indent,
            )

    def _title_teacher_block(self) -> None:
        """Преподаватель: метка → "<степень> <должность> ФИО"."""
        cfg = self._title
        if not cfg.teacher_name:
            return
        self._add_paragraph(left_indent=_RIGHT_BLOCK_INDENT,
                            align=WD_ALIGN_PARAGRAPH.LEFT)
        self._add_runs_paragraph(
            [
                StyledRun(cfg.teacher_label, underline=True),
                StyledRun(":"),
            ],
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=_RIGHT_BLOCK_INDENT,
        )
        parts = []
        if cfg.teacher_degree:
            parts.append(cfg.teacher_degree)
        if cfg.teacher_position:
            parts.append(cfg.teacher_position)
        parts.append(cfg.teacher_name)
        self._add_paragraph(
            " ".join(parts),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=_RIGHT_BLOCK_INDENT,
        )

    def _title_footer(self) -> None:
        """Город и год — плавающая borderless-таблица, прижатая к нижнему
        полю чтобы не уезжать на page 2 при длинной шапке/теме/ФИО."""
        cfg = self._title
        city = self._resolve("city")
        footer_lines = []
        if city:
            footer_lines.append(city)
        footer_lines.append(cfg.year)
        self._add_bottom_anchored_block(footer_lines)

    # --------------------------------------------------------
    # Публичный API: контент основного текста
    # --------------------------------------------------------

    def _enable_update_fields_on_open(self):
        """Ставит <w:updateFields w:val="true"/> в settings.xml. Word, Pages
        и LibreOffice при первом открытии файла обновят все поля (включая
        TOC и нумерацию). Идемпотентно: повторный r.toc() не дублирует."""
        settings = self._doc.settings.element
        existing = settings.find(qn("w:updateFields"))
        if existing is None:
            uf = OxmlElement("w:updateFields")
            uf.set(qn("w:val"), "true")
            settings.insert(0, uf)
        elif existing.get(qn("w:val")) != "true":
            existing.set(qn("w:val"), "true")

    def toc(self):
        """Вставляет автоматическое оглавление (поле Word TOC). Заголовок
        — `profile.toc_title` (по умолчанию «СОДЕРЖАНИЕ» в GOST_PROFILE,
        «ОГЛАВЛЕНИЕ» в ITMO_PROFILE).

        Также ставит флаг `updateFields=true` в settings.xml — Word/Pages
        при первом открытии файла предложит обновить поля (один клик
        «Yes» в диалоге), без ручного ПКМ по полю TOC.
        """
        self._enable_update_fields_on_open()
        self._add_paragraph(
            self._profile.toc_title,
            size=Pt(self._profile.heading_size_h1),
            bold=True,
            space_after=Pt(12),
        )

        p = self._doc.add_paragraph()
        run = p.add_run()
        _set_run_font(run)
        _append_word_field(
            run,
            ' TOC \\o "1-3" \\h \\z \\u ',
            placeholder=("Оглавление будет сформировано автоматически "
                         "при обновлении поля (ПКМ → «Обновить поле»)"),
            preserve_space=True,
        )

        self.page_break()

    def h1(self, text: str):
        """Заголовок 1 уровня. Поведение управляется профилем:
        h1_new_page (новая страница), h1_uppercase (ВЕРХНИЙ регистр).
        """
        if self._profile.h1_new_page and not self._just_broke_page:
            self.page_break()
        p = self._doc.add_paragraph(style="Heading 1")
        text = _sanitize_prose(text)
        text_to_render = text.upper() if self._profile.h1_uppercase else text
        run = p.add_run(text_to_render)
        _set_run_font(run, size=Pt(self._profile.heading_size_h1), bold=True)
        self._just_broke_page = False

    def h2(self, text: str):
        p = self._doc.add_paragraph(style="Heading 2")
        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run, size=Pt(self._profile.heading_size_h2), bold=True)

    def h3(self, text: str):
        p = self._doc.add_paragraph(style="Heading 3")
        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run, size=Pt(self._profile.heading_size_h3), bold=True)

    def text(self, text: str, *, bold=False, italic=False):
        self._add_paragraph(
            _sanitize_prose(text),
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=FIRST_LINE_INDENT,
            bold=bold, italic=italic,
        )

    def task(self, text: str):
        self._add_paragraph(
            _sanitize_prose(text),
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=FIRST_LINE_INDENT,
            bold=True,
            space_before=SPACE_BLOCK, space_after=SPACE_BLOCK,
        )

    def code(self, code: str):
        for line in code.split("\n"):
            p = self._make_paragraph(
                align=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=LINE_SPACING_CODE,
                first_line_indent=Pt(0),
                left_indent=Cm(0.5),
            )
            run = p.add_run(line if line else " ")
            run.font.name = "Courier New"
            run.font.size = Pt(11)
            _force_run_fonts(run._element.get_or_add_rPr(), "Courier New")

    def _resolve_figure_path(self, image_path: Union[str, Path]) -> Path:
        """Абсолютный путь — без изменений; относительный — от
        `self._paths.figures`. Бросает FileNotFoundError с обоими путями
        (input + resolved) если файл не существует.

        BC note: ранее относительные пути резолвились от os.getcwd(); теперь
        от <project>/docs/figures/. Реальные user-скрипты передавали
        абсолютные пути через Path(__file__).parent, поэтому wild-impact
        близок к нулю. Подробнее — CHANGELOG v0.7.0.
        """
        p = Path(image_path)
        if p.is_absolute():
            resolved = p
        else:
            resolved = self._paths.figures / p
        if not resolved.exists():
            raise FileNotFoundError(
                f"figure not found: input={image_path!r}, resolved={resolved}"
            )
        return resolved

    def figure(self, image, caption: str, *,
               width_cm: Optional[float] = None):
        """Вставляет рисунок с автоматической подписью.

        `image` — путь к PNG/JPEG (str/Path) ИЛИ Figure-объект из подключаемого
        модуля (r.plot.line(...), r.diagram(...)). Все варианты делят ОДИН
        сквозной счётчик «Рисунок N — …» (ГОСТ 7.32 §6.5).

        Ширина картинки **всегда** ограничивается печатной областью страницы
        (A4 минус левое и правое поля активного профиля — обычно ~17 см для
        ITMO и ~16.5 см для GOST). Если width_cm не задан, используется
        натуральный размер картинки (с клампом к печатной области, если она
        больше). Если width_cm задан — он уважается, но также клампится.
        Это предотвращает вылет крупных скриншотов за поля.

        Резолв пути:
        - абсолютный путь — без изменений.
        - относительный — резолвится от `<project>/docs/figures/`.
          `r.figure("load.png", ...)` → `<project>/docs/figures/load.png`.
          `r.figure("subdir/load.png", ...)` — тоже от `figures/`.
        Если файл не найден — FileNotFoundError с обоими путями (input + resolved).
        """
        from core_api import is_figure_like
        if is_figure_like(image):
            image_path = Path(image.render(
                dpi=300, max_width_cm=self._printable_cm()))
        else:
            image_path = self._resolve_figure_path(image)
        self._figure_counter += 1

        max_width_cm = self._printable_cm()

        image_path_str = str(image_path)
        if width_cm is None:
            # Прочитать натуральные размеры через python-docx (без PIL-зависимости)
            from docx.image.image import Image as _DocxImage
            img_meta = _DocxImage.from_file(image_path_str)
            natural_width_cm = img_meta.px_width / img_meta.horz_dpi * 2.54
            if natural_width_cm > max_width_cm:
                width_cm = max_width_cm
            # else: натуральный размер влезает, оставляем width_cm=None
        elif width_cm > max_width_cm:
            width_cm = max_width_cm

        p = self._make_paragraph(
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            space_before=SPACE_BLOCK,
        )
        run = p.add_run()
        if width_cm is not None:
            run.add_picture(image_path_str, width=Cm(width_cm))
        else:
            run.add_picture(image_path_str)

        self._add_paragraph(
            f"Рисунок {self._figure_counter} — {_sanitize_prose(caption)}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            space_after=SPACE_BLOCK,
        )
        return self._figure_counter

    def formula(self, latex: str, *, where: Optional[str] = None) -> int:
        """Вставляет формулу из LaTeX как нативное Word-уравнение (OMML).

        Возвращает номер формулы — пригодится для ссылок:
            f1 = r.formula(r"E = mc^2")
            r.text(f"По формуле ({f1}) видно ...")

        Параметр `where` — пояснения переменных, идут отдельным абзацем
        под формулой:
            r.formula(
                r"\\frac{a + b}{c}",
                where="a, b — слагаемые, c — делитель",
            )

        Сама строка `latex` не санируется (LaTeX-синтаксис не должен
        искажаться); `where` проходит через _sanitize_prose как обычная
        проза. Зависит от пакета `latex2mathml` — он ставится автоматически
        через scripts/ensure_env.py.

        Форматирование по ГОСТ 7.32: формула центрирована, номер «(N)»
        прижат к правому краю печатной области на той же строке.

        BC-shim: реализация вынесена в модуль gost_report_math (r.math.formula).
        Поведение и раскладка идентичны прежним.
        """
        return self.math.formula(latex, where=where)

    def table(self, rows: Sequence[Sequence[str]], caption: str = "",
              *, has_header: bool = True):
        if not rows:
            return
        self._table_counter += 1

        if caption:
            self._add_paragraph(
                f"Таблица {self._table_counter} — {_sanitize_prose(caption)}",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Pt(0),
                space_before=SPACE_BLOCK,
            )

        n_cols = max(len(r) for r in rows)
        table = self._doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Table Grid"
        for i, row_data in enumerate(rows):
            for j in range(n_cols):
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = LINE_SPACING_TABLE
                p.paragraph_format.first_line_indent = Pt(0)
                value = row_data[j] if j < len(row_data) else ""
                run = p.add_run(_sanitize_prose(value))
                _set_run_font(run, bold=(has_header and i == 0))
        return self._table_counter

    def _create_independent_num(self, abstract_num_id: str) -> int:
        numbering = self._doc.part.numbering_part.element
        num_id = self._next_num_id
        self._next_num_id += 1

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), abstract_num_id)
        num.append(abs_ref)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)

        numbering.append(num)
        return num_id

    def _find_abstract_num_for_style(self, style_name: str) -> Optional[str]:
        try:
            style = self._doc.styles[style_name]
        except KeyError:
            return None
        style_pPr = style.element.find(qn("w:pPr"))
        if style_pPr is None:
            return None
        numPr = style_pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        numId_el = numPr.find(qn("w:numId"))
        if numId_el is None:
            return None
        style_num_id = numId_el.get(qn("w:val"))

        numbering = self._doc.part.numbering_part.element
        for num in numbering.findall(qn("w:num")):
            if num.get(qn("w:numId")) == style_num_id:
                abs_ref = num.find(qn("w:abstractNumId"))
                if abs_ref is not None:
                    return abs_ref.get(qn("w:val"))
        return None

    def _add_list_paragraph(self, text: str, num_id: int):
        # align и space_before должны прийти из стиля списка
        # (`List Number` / `List Bullet`). Без _INHERIT мы бы навязали
        # alignment=CENTER и w:before="0" в каждом list-параграфе.
        p = self._make_paragraph(
            align=_INHERIT,
            space_before=_INHERIT,
        )

        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(num_id))
        numPr.append(ilvl)
        numPr.append(numId_el)
        pPr.append(numPr)

        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run)

    def _ensure_numbering_part(self, style_name: str) -> None:
        """Триггерит создание numbering.xml части документа, если её ещё нет.
        python-docx создаёт numbering_part лениво, при первом использовании
        стиля списка. Бросаем туда-сюда пустой параграф со стилем чтобы
        гарантировать наличие numbering_part до _find_abstract_num_for_style."""
        try:
            _ = self._doc.part.numbering_part
        except (AttributeError, KeyError):
            tmp = self._doc.add_paragraph(style=style_name)
            tmp._element.getparent().remove(tmp._element)

    def _list(self, items, *, style_name: str) -> None:
        if isinstance(items, str):
            items = [items]
        if not items:
            return
        self._ensure_numbering_part(style_name)
        abstract_id = self._find_abstract_num_for_style(style_name)
        if abstract_id is None:
            for item in items:
                p = self._doc.add_paragraph(style=style_name)
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING_BODY
                pf.space_after = Pt(0)
                run = p.add_run(_sanitize_prose(item))
                _set_run_font(run)
            return
        num_id = self._create_independent_num(abstract_id)
        for item in items:
            self._add_list_paragraph(item, num_id)

    def numbered(self, items):
        self._list(items, style_name="List Number")

    def bullet(self, items):
        self._list(items, style_name="List Bullet")

    def page_break(self):
        p = self._doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        self._just_broke_page = True

    def save(self, path: Optional[Union[str, Path]] = None) -> Path:
        """Сохранить .docx. Если path не задан — `<project>/docs/report.docx`.
        Относительный путь — от `<project>/docs/`. Абсолютный — без изменений.
        Создаёт parent-директории при необходимости. Возвращает абсолютный Path.

        После записи файл валидируется автоматически через `validate.py`.
        Если найдены нарушения уровня FAIL — поднимается `GostValidationError`.
        Тёплые предупреждения (heuristic) печатаются в stderr и не блокируют.
        Если validate.py недоступен (битый venv) — save() работает как раньше,
        отсутствие валидации лучше падения библиотеки.
        """
        if path is None:
            target = self._paths.out
        else:
            p = Path(path)
            target = p if p.is_absolute() else self._paths.docs / p
        target.parent.mkdir(parents=True, exist_ok=True)
        _strip_python_docx_fingerprints(self._doc)
        self._doc.save(str(target))
        _strip_thumbnail_part(target)
        target = target.resolve()

        # Картинки уже встроены байтами в .docx — tmp-PNG модулей больше не нужны.
        self._cleanup()

        try:
            import validate as _validate
        except ImportError:
            sys.stderr.write(
                "gost-report: validate.py недоступен, проверка пропущена.\n"
            )
            return target

        try:
            profile_name = type(self._profile).__name__
            for var_name, value in globals().items():
                if value is self._profile and var_name.endswith("_PROFILE"):
                    profile_name = var_name
                    break
            _validate.write_sentinel(target, profile_name=profile_name)
            violations = _validate.validate_docx(target)
        except Exception as e:
            sys.stderr.write(
                f"gost-report: валидатор упал ({e}); проверка пропущена.\n"
            )
            return target

        warns = [v for v in violations if v.tier == _validate.TIER_WARN]
        for v in warns:
            sys.stderr.write("gost-report warn: " + v.format().lstrip() + "\n")

        if _validate.has_failures(violations):
            raise GostValidationError(
                _validate.format_error(target, violations)
            )
        return target

    @property
    def doc(self) -> _Document:
        return self._doc

    @property
    def profile(self) -> UniversityProfile:
        return self._profile
