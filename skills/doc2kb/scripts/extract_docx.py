#!/usr/bin/env python3
"""
extract_docx.py — Phase 4 extractor for DOCX files.

CLI:
    extract_docx.py <input_docx> <output_md>
                    [--doc-id doc-NNN]
                    [--source-rel relative/path/in/corpus.docx]
                    [--assets-dir <abs_path>]
                    [--assets-rel <rel_prefix_from_md>]
                    [--no-extract-images]
                    [--force-mammoth]

Effect:
    Writes <output_md> with YAML frontmatter + Markdown body. Stdout receives
    a single JSON line summarizing the result.

Routing:
    - If the document contains OOXML math (`<m:oMath>` / `<m:oMathPara>`)
      AND `pandoc` is on PATH, the body is produced by pandoc (math survives
      as `$...$` / `$$...$$` LaTeX). Mammoth's HTML/markdown writers both
      silently drop oMath elements, which destroys the math content the
      second-session agent will be asked about.
    - Otherwise (no math, or pandoc unavailable) the body is produced by
      mammoth → HTML → markdownify, which preserves headings/lists/tables
      better than pandoc for prose-only docs.
    - `--force-mammoth` bypasses the pandoc route even when math is present
      (regression-test escape hatch).

Image handling:
    - By default, inline images are extracted via save_image_safe() to the
      kb's `assets/` directory and referenced from the body as proper
      Markdown image links. Mammoth feeds each image to a custom handler
      that writes the bytes and returns the relative path as `src`. The
      pandoc route uses `--extract-media` and we then rewrite the paths
      through save_image_safe so the 2000-px cap still applies.
    - `--no-extract-images` reverts to the old behaviour where images
      become `![image N]()` placeholders (no `src`).

Pipeline (mammoth route):
    1. mammoth.convert_to_html(input) — semantic conversion preserving
       headings, lists, tables, inline images, footnotes.
    2. markdownify(html) — HTML → Markdown.
    3. python-docx scout for `inline_images`, `tables`, `paragraphs`,
       `has_equations`, `has_tracked_changes` (already gathered by scout
       but we re-derive here so the extractor is callable standalone).

Why HTML→Markdown instead of mammoth.convert_to_markdown:
    mammoth's built-in markdown writer drops tables. The HTML route keeps
    them, even if mammoth's HTML for tables is plain (no header markers
    beyond <th> — but markdownify renders them as Markdown pipe tables).
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from _common import (  # noqa: E402
    add_assets_args,
    clean_whitespace,
    count_tokens,
    emit_failure,
    emit_success,
    resolve_assets_dir,
    sanitize_heading,
    save_image_safe,
    sha256_of,
    today_iso,
    tool_version_string,
    validate_source_rel,
    write_md,
)


EXTRACTOR_MAMMOTH = "mammoth+markdownify"
EXTRACTOR_PANDOC = "pandoc"
# Below this body length on a non-empty docx, flag a warning.
MIN_BODY_CHARS = 50
# Hard cap on pandoc invocation; pandoc on a typical lab docx finishes in
# <500ms, so 60s is generous and still catches a wedged process.
PANDOC_TIMEOUT_SEC = 60


def _try_import():
    try:
        import mammoth  # type: ignore
        from markdownify import markdownify as md_from_html  # type: ignore
        from docx import Document  # type: ignore
        return mammoth, md_from_html, Document
    except Exception as e:
        emit_failure(f"required libraries unavailable: {e}")
        sys.exit(1)


def _scout_docx(path: Path) -> dict:
    """Light re-scan to populate frontmatter — duplicates scout_corpus.py
    logic but keeps the extractor invocable without scout output."""
    from docx import Document  # type: ignore
    info = {
        "paragraphs": 0,
        "inline_images": 0,
        "has_tables": False,
        "has_equations": False,
        "has_tracked_changes": False,
    }
    try:
        doc = Document(str(path))
        info["paragraphs"] = len(doc.paragraphs)
        info["inline_images"] = len(doc.inline_shapes)
        info["has_tables"] = len(doc.tables) > 0
        xml = doc.element.xml
        info["has_equations"] = "<m:oMath" in xml or "<m:oMathPara" in xml
        info["has_tracked_changes"] = "w:ins" in xml or "w:del" in xml
    except Exception:
        pass
    return info


def _extract_via_mammoth(
    input_docx: Path,
    doc_id: str = "doc-000",
    assets_dir: Path | None = None,
    assets_rel_prefix: str = "../assets",
    extract_images: bool = True,
) -> tuple[str, list[str], list[str]]:
    """Mammoth → HTML → Markdown route. Returns (body, warnings, assets).

    When `extract_images=True` and `assets_dir` is provided, the mammoth
    image handler funnels each picture through save_image_safe() so the
    2000-px cap and decorative-pixel filter apply uniformly with the PDF
    and PPTX extractors. Otherwise mammoth's old placeholder-only behaviour
    is preserved (alt text + count, empty `src`).
    """
    mammoth, md_from_html, _Document = _try_import()
    warnings: list[str] = []
    assets: list[str] = []
    do_extract = extract_images and assets_dir is not None
    if do_extract:
        try:
            assets_dir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            warnings.append(
                f"cannot create assets dir {assets_dir}: {e} — falling back to placeholder-only images"
            )
            do_extract = False
    dedup_cache: dict = {}
    img_idx = {"n": 0}

    def _image_to_asset(image):
        img_idx["n"] += 1
        n = img_idx["n"]
        alt = (image.alt_text or "").strip()
        # Read mammoth's bytes via the context manager it gives us.
        try:
            ext = "png"
            ct = (image.content_type or "").lower()
            # content_type looks like "image/png" or "image/jpeg".
            if "/" in ct:
                ext = ct.split("/", 1)[1] or "png"
            with image.open() as fh:
                blob = fh.read()
        except Exception as e:
            warnings.append(f"image {n}: cannot read bytes ({e})")
            label = f"image {n}: {alt}" if alt else f"image {n}"
            return {"src": "", "alt": label[:120]}

        filename = f"{doc_id}-img{n:03d}.{ext}"
        target = assets_dir / filename
        result = save_image_safe(
            blob, target, dedup_cache=dedup_cache, source_ext=ext,
        )
        if not result:
            if result.reason == "skipped_format":
                warnings.append(
                    f"image {n}: skipped non-viewable format ({ext}) — "
                    "vector/metafile images cannot be rendered by the vision encoder"
                )
            elif result.reason == "skipped_corrupt":
                warnings.append(f"image {n}: corrupt image bytes — skipped")
            label = f"image {n}: {alt}" if alt else f"image {n}"
            return {"src": "", "alt": label[:120]}

        rel = f"{assets_rel_prefix}/{result.path.name}"
        if rel not in assets:
            assets.append(rel)
        label = alt or f"image {n}"
        return {"src": rel, "alt": label[:120]}

    def _image_placeholder(image):
        # Legacy behaviour when image extraction is disabled.
        img_idx["n"] += 1
        n = img_idx["n"]
        alt = (image.alt_text or "").strip()
        label = f"image {n}: {alt}" if alt else f"image {n}"
        return {"src": "", "alt": label[:120]}

    image_callback = _image_to_asset if do_extract else _image_placeholder
    image_handler = mammoth.images.img_element(image_callback)
    try:
        with input_docx.open("rb") as fh:
            result = mammoth.convert_to_html(fh, convert_image=image_handler)
    except Exception as e:
        raise RuntimeError(f"mammoth conversion failed: {e}") from e

    html = result.value or ""
    # mammoth's messages are tuples of (type, message); promote warnings.
    for m in result.messages or []:
        try:
            mt = m.type if hasattr(m, "type") else m[0]
            mm = m.message if hasattr(m, "message") else m[1]
        except Exception:
            mt = "warning"
            mm = str(m)
        # mammoth reports lots of style-mapping notices that are noise here.
        if mt == "warning" and "style" not in str(mm).lower():
            warnings.append(f"mammoth: {mm}"[:200])

    # HTML → Markdown. heading_style=ATX renders "# heading" instead of underline.
    md = md_from_html(html, heading_style="ATX", bullets="-",
                      strip=["style", "script"])
    body = clean_whitespace(md)
    return body, warnings, assets


def _extract_via_pandoc(
    input_docx: Path,
    doc_id: str = "doc-000",
    assets_dir: Path | None = None,
    assets_rel_prefix: str = "../assets",
    extract_images: bool = True,
) -> tuple[str, list[str], list[str]]:
    """Pandoc route — used when the source contains OOXML math that mammoth
    would silently drop. Pandoc emits math as `$...$` / `$$...$$` LaTeX,
    which the LLM can read and render in any downstream context.

    Pandoc also handles tables, lists, footnotes, and headings competently;
    the only reason we still default to mammoth for prose-only docs is that
    mammoth produces slightly cleaner list nesting and avoids pandoc's
    blockquote-wrapping of continuation paragraphs inside numbered lists.

    When image extraction is enabled, pandoc writes images to a temporary
    directory via `--extract-media`; we post-process those into
    `<assets_dir>/<doc_id>-img<NNN>.<ext>` through save_image_safe(), and
    rewrite the body to point at the cleaned-up paths.

    Returns (body, warnings, assets)."""
    warnings: list[str] = []
    assets: list[str] = []
    do_extract = extract_images and assets_dir is not None
    cmd = [
        "pandoc",
        "-f", "docx",
        "-t", "markdown",
        "--wrap=none",
        str(input_docx),
    ]

    tmp_media: Path | None = None
    if do_extract:
        try:
            assets_dir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            warnings.append(
                f"cannot create assets dir {assets_dir}: {e} — pandoc images will be dropped"
            )
            do_extract = False
    if do_extract:
        tmp_media = Path(tempfile.mkdtemp(prefix="doc2kb-pandoc-media-"))
        cmd.insert(1, "--extract-media")
        cmd.insert(2, str(tmp_media))

    try:
        proc = subprocess.run(
            cmd,
            capture_output=True,
            timeout=PANDOC_TIMEOUT_SEC,
            check=False,
        )
    except subprocess.TimeoutExpired:
        if tmp_media is not None:
            shutil.rmtree(tmp_media, ignore_errors=True)
        raise RuntimeError(f"pandoc timed out after {PANDOC_TIMEOUT_SEC}s")
    except FileNotFoundError:
        if tmp_media is not None:
            shutil.rmtree(tmp_media, ignore_errors=True)
        raise RuntimeError("pandoc binary not found on PATH")
    if proc.returncode != 0:
        stderr = (proc.stderr or b"").decode("utf-8", "replace").strip()
        if tmp_media is not None:
            shutil.rmtree(tmp_media, ignore_errors=True)
        raise RuntimeError(
            f"pandoc exited with code {proc.returncode}: {stderr[:200]}"
        )
    body = proc.stdout.decode("utf-8", "replace")
    body = clean_whitespace(body)
    if proc.stderr:
        stderr = proc.stderr.decode("utf-8", "replace").strip()
        for ln in stderr.splitlines():
            ln = ln.strip()
            if ln and ln.lower().startswith(("warning", "[warning]")):
                warnings.append(f"pandoc: {ln}"[:200])

    # Post-process pandoc's extracted media. Pandoc names files like
    # `<tmp_media>/media/image1.png`; we feed each through save_image_safe
    # (which enforces the 2000-px cap) and rewrite the body to point at
    # the canonical kb path.
    if tmp_media is not None:
        try:
            body, pandoc_assets, pandoc_warnings = _rewrite_pandoc_media(
                body, tmp_media, assets_dir, assets_rel_prefix, doc_id,
            )
            assets.extend(pandoc_assets)
            warnings.extend(pandoc_warnings)
        finally:
            shutil.rmtree(tmp_media, ignore_errors=True)

    return body, warnings, assets


_PANDOC_IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def _rewrite_pandoc_media(
    body: str,
    tmp_media: Path,
    assets_dir: Path,
    assets_rel_prefix: str,
    doc_id: str,
) -> tuple[str, list[str], list[str]]:
    """Walk the pandoc-extracted media tree, push each image through
    save_image_safe(), and rewrite every Markdown image link in `body`
    from the temp path to the canonical `<assets_rel_prefix>/<filename>`.
    Pandoc emits paths with forward slashes on every platform, so a
    string-based rewrite is sufficient — no regex over POSIX/Windows
    separators.
    """
    warnings: list[str] = []
    assets: list[str] = []
    dedup_cache: dict = {}
    # path_in_body (string) → new relative path (string)
    rewrite: dict[str, str] = {}
    idx = 0
    for match in _PANDOC_IMG_RE.finditer(body):
        link = match.group(2).strip()
        # Pandoc uses POSIX paths even on Windows; strip optional `<...>`
        # angle brackets pandoc sometimes wraps around paths with spaces.
        if link.startswith("<") and link.endswith(">"):
            link = link[1:-1]
        src = (tmp_media / link.replace("\\", "/")).resolve() \
            if not Path(link).is_absolute() else Path(link).resolve()
        # Guardrail: never read outside tmp_media — defends against an
        # untrusted .docx coaxing pandoc into emitting a `..` traversal.
        try:
            src.relative_to(tmp_media.resolve())
        except ValueError:
            warnings.append(
                f"pandoc emitted image path outside media dir, skipping: {link}"
            )
            rewrite[match.group(2)] = ""
            continue
        if not src.is_file():
            continue
        if link in rewrite:
            continue
        idx += 1
        ext = src.suffix.lstrip(".").lower() or "png"
        filename = f"{doc_id}-img{idx:03d}.{ext}"
        target = assets_dir / filename
        try:
            blob = src.read_bytes()
        except OSError as e:
            warnings.append(f"pandoc image {src.name}: cannot read ({e})")
            rewrite[match.group(2)] = ""
            continue
        result = save_image_safe(
            blob, target, dedup_cache=dedup_cache, source_ext=ext,
        )
        if not result:
            if result.reason == "skipped_format":
                warnings.append(
                    f"pandoc image {src.name}: skipped non-viewable format ({ext})"
                )
            elif result.reason == "skipped_corrupt":
                warnings.append(
                    f"pandoc image {src.name}: corrupt image bytes — skipped"
                )
            rewrite[match.group(2)] = ""
            continue
        rel = f"{assets_rel_prefix}/{result.path.name}"
        rewrite[match.group(2)] = rel
        if rel not in assets:
            assets.append(rel)

    if not rewrite:
        return body, assets, warnings

    def _repl(m):
        new = rewrite.get(m.group(2))
        if new is None:
            return m.group(0)
        if new == "":
            # Drop the image element entirely when the source was unusable.
            return f"![{m.group(1)}]()"
        return f"![{m.group(1)}]({new})"

    new_body = _PANDOC_IMG_RE.sub(_repl, body)
    return new_body, assets, warnings


def extract(input_docx: Path,
            doc_id: str = "doc-000",
            assets_dir: Path | None = None,
            assets_rel_prefix: str = "../assets",
            extract_images: bool = True,
            force_mammoth: bool = False) -> tuple[str, dict, list[str], str]:
    """Returns (markdown_body, frontmatter_extras, warnings, extractor_name).

    Routes math-bearing documents through pandoc (when available) so OOXML
    math survives as LaTeX instead of being silently dropped. See the
    module docstring for the full routing table."""
    extras = _scout_docx(input_docx)
    pandoc_available = shutil.which("pandoc") is not None
    use_pandoc = (
        extras.get("has_equations") and pandoc_available and not force_mammoth
    )
    warnings: list[str] = []
    assets: list[str] = []

    if use_pandoc:
        try:
            body, w, a = _extract_via_pandoc(
                input_docx,
                doc_id=doc_id,
                assets_dir=assets_dir,
                assets_rel_prefix=assets_rel_prefix,
                extract_images=extract_images,
            )
            warnings.extend(w)
            assets.extend(a)
            extractor_name = EXTRACTOR_PANDOC
        except RuntimeError as e:
            # Pandoc failure → fall back to mammoth with a loud warning so
            # the operator knows math was dropped.
            warnings.append(
                f"pandoc route failed ({e}); falling back to mammoth — "
                "OOXML math elements will be dropped from the output"
            )
            body, w, a = _extract_via_mammoth(
                input_docx,
                doc_id=doc_id,
                assets_dir=assets_dir,
                assets_rel_prefix=assets_rel_prefix,
                extract_images=extract_images,
            )
            warnings.extend(w)
            assets.extend(a)
            extractor_name = EXTRACTOR_MAMMOTH
    else:
        body, w, a = _extract_via_mammoth(
            input_docx,
            doc_id=doc_id,
            assets_dir=assets_dir,
            assets_rel_prefix=assets_rel_prefix,
            extract_images=extract_images,
        )
        warnings.extend(w)
        assets.extend(a)
        extractor_name = EXTRACTOR_MAMMOTH
        if extras.get("has_equations") and not pandoc_available:
            warnings.append(
                "has_equations=true but pandoc is not on PATH — OOXML math "
                "elements (m:oMath) were dropped by mammoth. Install pandoc "
                "(`brew install pandoc` / `apt install pandoc`) and re-run "
                "to preserve formulas as LaTeX"
            )

    if len(body.strip()) < MIN_BODY_CHARS:
        warnings.append(f"extracted body is unusually short ({len(body)} chars)")

    headings = []
    for line in body.split("\n"):
        if line.startswith("# ") and len(line) < 200:
            sanitized = sanitize_heading(line[2:])
            if sanitized:
                headings.append(sanitized)
            if len(headings) >= 10:
                break
    extras["headings"] = headings
    if assets:
        extras["assets"] = assets

    return body, extras, warnings, extractor_name


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Extract DOCX → Markdown. Routes through pandoc when the "
                    "source contains OOXML math; otherwise via mammoth + markdownify."
    )
    ap.add_argument("input_docx")
    ap.add_argument("output_md")
    ap.add_argument("--doc-id", default="doc-000")
    ap.add_argument("--source-rel", default=None)
    add_assets_args(ap)
    ap.add_argument("--force-mammoth", action="store_true",
                    help="Bypass the pandoc-for-math route even when math is present.")
    args = ap.parse_args()

    in_path = Path(args.input_docx).expanduser().resolve()
    out_path = Path(args.output_md).expanduser().resolve()
    source_rel = args.source_rel or in_path.name
    try:
        source_rel = validate_source_rel(source_rel)
    except ValueError as e:
        emit_failure(f"invalid --source-rel: {e}")
        return 1

    if not in_path.is_file():
        emit_failure(f"input not found: {in_path}")
        return 1

    assets_dir: Path | None = resolve_assets_dir(args, out_path)

    try:
        body, extras, warnings, extractor_name = extract(
            in_path,
            doc_id=args.doc_id,
            assets_dir=assets_dir,
            assets_rel_prefix=args.assets_rel,
            extract_images=not args.no_extract_images,
            force_mammoth=args.force_mammoth,
        )
    except Exception as e:
        emit_failure(f"extraction failed: {e}", extra={"input": str(in_path)})
        return 1

    fm = {
        "id": args.doc_id,
        "source": source_rel,
        "source_type": "docx",
        "source_sha256": sha256_of(in_path),
        "extraction_method": f"{extractor_name}@{tool_version_string()}",
        "extraction_date": today_iso(),
        "paragraphs": extras.get("paragraphs", 0),
        "inline_images": extras.get("inline_images", 0),
        "has_tables": extras.get("has_tables", False),
        "has_equations": extras.get("has_equations", False),
        "has_tracked_changes": extras.get("has_tracked_changes", False),
        "headings": extras.get("headings", []),
        "tokens_estimated": count_tokens(body),
        "warnings": warnings,
    }
    if extras.get("assets"):
        fm["assets"] = extras["assets"]
    write_md(out_path, fm, body)
    emit_success(out_path, body, warnings, extra={
        "inline_images": extras.get("inline_images", 0),
        "has_tables": extras.get("has_tables", False),
        "extractor": extractor_name,
        "assets_extracted": len(extras.get("assets", [])),
    })
    return 0


if __name__ == "__main__":
    sys.exit(main())
